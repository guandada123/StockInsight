"""生成小白也能看懂的个股深度分析 Word 报告"""
import sys, os
from datetime import datetime


def generate(code: str):
    from stock_analyzer.cache import cached_kline, cached_fundamentals, cached_fund_flow
    from stock_analyzer.analysis import full_technical_analysis, get_technical_summary, calc_support_resistance, calc_stop_levels
    from stock_analyzer.quant import composite_quant_score, calc_risk_metrics
    from stock_analyzer.short_term import short_term_score, calc_combo_signals, calc_multi_timeframe_resonance
    from stock_analyzer.fetcher import sina_real_time, get_market_overview
    from stock_analyzer.sector_info import get_stock_sector_full
    from stock_analyzer.ml_predict import _cached_predict_ensemble
    from stock_analyzer.advanced import macro_market_signal
    from stock_analyzer.nl_report import generate_bull_bear_debate
    from stock_analyzer.backtest import compare_strategies, DEFAULT_COMPARE_STRATEGIES
    from stock_analyzer.patterns import generate_kline_interpretation_with_today
    from stock_analyzer.psychology import analyze_manipulator_intention

    print(f"正在分析 {code} ...")
    kline = full_technical_analysis(cached_kline(code, days=365))
    funds = cached_fundamentals(code)
    tech = get_technical_summary(kline)
    sr = calc_support_resistance(kline)
    price = float(kline["收盘"].iloc[-1])
    atr = float(kline.iloc[-1].get("ATR", price * 0.03))
    stop = calc_stop_levels(price, atr, float(sr.get("支撑位", [price * 0.9])[0]), float(sr.get("压力位", [price * 1.1])[0]))
    quant = composite_quant_score(kline, funds)
    risk = calc_risk_metrics(kline)
    st = short_term_score(kline, code)
    combo = calc_combo_signals(kline)
    mr = calc_multi_timeframe_resonance(code)
    rt = sina_real_time([code])
    # Note: During trading hours (9:30-15:00 Mon-Fri), real-time price differs from cached close
    # The analysis uses cached K-line for technical indicators and real-time price for current value
    # On weekends (Sat-Sun), data reflects last trading day (Friday)
    info = rt.get(code, {})
    name = info.get("名称", code)
    market = get_market_overview()
    sector = get_stock_sector_full(code)
    ml = _cached_predict_ensemble(kline, funds)
    n5 = round((kline["收盘"].iloc[-1] / kline["收盘"].iloc[-6] - 1) * 100, 2)
    n20 = round((kline["收盘"].iloc[-1] / kline["收盘"].iloc[-21] - 1) * 100, 2)
    n60 = round((kline["收盘"].iloc[-1] / kline["收盘"].iloc[-61] - 1) * 100, 2)
    debate = generate_bull_bear_debate({
        "quant_score": quant.get("composite_score", 50),
        "technical": {"macd_signal": tech.get("macd_signal", ""), "kdj_signal": tech.get("kdj_signal", ""),
            "rsi": tech.get("rsi_value", 50), "near5d": n5, "near20d": n20,
            "ma_status": tech.get("均线", ""), "resistance": sr.get("压力位", []), "price": price, "pe": 0},
        "fund_flow": {"direction": "无数据"},
        "ai_prediction": {"direction": ml.get("ensemble_direction", "?"), "confidence": ml.get("ensemble_confidence", 50)}})
    try: ff = cached_fund_flow(code, days=5); total_flow = round(ff["主力净流入-净额"].sum() / 1e8, 2)
    except: total_flow = 0
    try: bt = compare_strategies(kline, DEFAULT_COMPARE_STRATEGIES, 100000)
    except: bt = {}
    try: macro = macro_market_signal()
    except: macro = {}

    # ── 综合评分等级 ──
    qs = quant.get("composite_score", 50)
    st_sc = st.get("短线评分", 50) if isinstance(st, dict) else 50
    combo_sig = combo.get("信号", "?")
    combo_str = combo.get("强度", 0)
    mr_str = mr.get("共振强度", 0)
    rsi_val = tech.get("rsi_value", 50)
    macd_sig = tech.get("macd_signal", "")
    ml_dir = ml.get("ensemble_direction", "?")
    ml_conf = ml.get("ensemble_confidence", 50)
    ml_agree = ml.get("agreement", "?")
    roe = funds.get("ROE", 0) if isinstance(funds, dict) else 0
    pe = funds.get("市盈率", "N/A") if isinstance(funds, dict) else "N/A"
    pb = funds.get("市净率", "N/A") if isinstance(funds, dict) else "N/A"
    eps = funds.get("每股收益", "N/A") if isinstance(funds, dict) else "N/A"

    # 综合判断
    buy_signals = 0
    if qs >= 60: buy_signals += 1
    if st_sc >= 60: buy_signals += 1
    if combo_str >= 3: buy_signals += 1
    if ml_dir == "看涨" and ml_conf >= 55: buy_signals += 1
    if rsi_val < 35: buy_signals += 1
    if n5 > 0 and n20 < 10: buy_signals += 1

    if buy_signals >= 4:
        risk_level = "低风险 ✅"
        risk_color = "22c55e"
        verdict = "多项指标共振看多，适合买入"
    elif buy_signals >= 2:
        risk_level = "中等风险 ⚠️"
        risk_color = "f59e0b"
        verdict = "部分指标向好，可轻仓试探"
    else:
        risk_level = "高风险 ❌"
        risk_color = "ef4444"
        verdict = "多指标偏空，建议观望或减仓"

    buy_low = round(price - atr * 0.5, 2)
    buy_high = round(price + atr * 0.3, 2)

    # ═══════════════════════════════════
    # 生成 Word
    # ═══════════════════════════════════
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)

    def add_colored_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    def add_kpi_table(headers, rows):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]; cell.text = h
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(10); run.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]; cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs: run.font.size = Pt(10)
        doc.add_paragraph()
        return table

    # ═══ 封面 ═══
    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f'{name} 深度分析报告'); r.font.size = Pt(28); r.bold = True; r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f'股票代码: {code}  |  行业: {sector}').font.size = Pt(14)

    doc.add_paragraph()
    rl = doc.add_paragraph(); rl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rl.add_run(f'综合风险评级: {risk_level}'); rr.font.size = Pt(18); rr.bold = True
    rr.font.color.rgb = RGBColor(int(risk_color[0:2], 16), int(risk_color[2:4], 16), int(risk_color[4:6], 16))

    doc.add_paragraph()
    vp = doc.add_paragraph(); vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vp.add_run(f'📊 {verdict}').font.size = Pt(13)
    doc.add_paragraph(f'报告日期: {datetime.now().strftime("%Y-%m-%d")} (K线数据至 {str(kline.iloc[-1][chr(26085)+chr(26399)])[:10]}, 实时价 {price:.2f})')
    doc.add_page_break()

    # ═══ 一、一句话总结 ═══
    add_colored_heading('一、一句话总结', level=1)
    doc.add_paragraph(f'{name}（{code}）当前股价 {price} 元，近5日涨跌幅 {n5:+.1f}%。'
                      f'技术面{macd_sig}，RSI={rsi_val:.0f}，AI预测{ml_dir}（置信度{ml_conf:.0f}%）。'
                      f'综合评分 {qs:.0f} 分，属于"{quant.get("rating", "")}"级别。'
                      f'建议操作: {verdict}')
    doc.add_paragraph()

    # ═══ 二、大盘环境（今天市场整体怎么样）════
    add_colored_heading('二、大盘环境（今天市场整体怎么样）', level=1)
    doc.add_paragraph('大盘指数反映了整个股市的涨跌情况。如果大盘都在跌，个股也很难独善其身。')
    rows = []
    for ci, cn in [("000001", "上证指数（大盘股）"), ("399001", "深证成指（中小盘）"), ("399006", "创业板指（成长股）"), ("000688", "科创50（科技股）")]:
        idx = market.get(ci, {})
        p = float(idx.get("最新价", 0) or 0)
        c = float(idx.get("涨跌幅", 0) or 0)
        color = "🔴" if c < -1 else ("🟢" if c > 1 else "🟡")
        rows.append([cn, f"{p:.2f}", f"{c:+.2f}%", color + (" 跌" if c < 0 else " 涨")])
    add_kpi_table(["指数名称", "点位", "涨跌幅", "状态"], rows)
    doc.add_paragraph('👉 解读：红色表示市场在下跌，绿色在上涨。大盘全绿时买入要格外谨慎。')

    # ═══ 三、技术面（股票走势如何）════
    add_colored_heading('三、技术面分析（股票走势如何）', level=1)

    doc.add_paragraph('技术面通过价格和成交量的历史数据，判断股票当前处于什么状态。', style='Intense Quote')
    add_kpi_table(["指标", "当前值", "什么意思", "信号"],
                  [[f"最新价", f"{price:.2f} 元", "股票的当前市场价格", "—"],
                   [f"今日区间", f"{float(info.get('今开', 0)):.2f} - {float(info.get('最高', 0)):.2f}", f"今天最高到最低的范围", "—"],
                   [f"近5日涨跌", f"{n5:+.1f}%", "最近一周的涨跌。涨太多追高风险大，跌太多可能反弹", "⚠️ 追高" if n5 > 12 else ("✅ 温和" if n5 > 0 else "🟢 超跌反弹机会")],
                   [f"近20日涨跌", f"{n20:+.1f}%", "最近一个月的涨跌。超过30%说明涨太多了", "🔴 过热" if n20 > 30 else ("✅ 正常" if n20 > -10 else "🟢 超跌")],
                   [f"RSI", f"{rsi_val:.0f}", "强弱指标。>70太贵(超买)，<30太便宜(超卖)", f"{'🔴 超买' if rsi_val > 70 else ('🟢 超卖机会' if rsi_val < 30 else '✅ 正常')}"],
                   [f"MACD", macd_sig, "趋势方向。多头=上涨趋势，空头=下跌趋势", "✅ 向好" if "多头" in macd_sig else "⚠️ 偏弱"]])

    doc.add_paragraph()
    doc.add_paragraph(
        f'📌 关键价位：支撑位 {[round(float(x), 2) for x in sr.get("支撑位", [price * 0.9])[:2]]}（跌到这里可能止跌反弹）\n'
        f'   压力位 {[round(float(x), 2) for x in sr.get("压力位", [price * 1.1])[:2]]}（涨到这里可能遇阻回落）\n'
        f'   ATR(平均真实波幅): {atr:.2f} 元 — 数值越大说明股价波动越剧烈')


    # K-line patterns
    doc.add_page_break()
    doc.add_heading('K线形态分析', level=1)
    doc.add_paragraph('K线形态反映市场情绪，以下是近期出现的典型形态：', style='Intense Quote')
    try:
        interp = generate_kline_interpretation_with_today(kline)
        trend = interp.get('trend_phase', '')
        recent = interp.get('recent_patterns', [])
        summary = interp.get('summary', '')
        key_obs = interp.get('key_observation', '')

        doc.add_paragraph('当前趋势阶段：' + str(trend))

        if recent:
            for p in recent:
                tag = '看涨' if p.get('type') == 'bullish' else '看跌'
                doc.add_paragraph('  ' + tag + ': ' + p['name'] + ' (' + p['date'] + ')')
                doc.add_paragraph('    ' + p['description'][:120])
        else:
            doc.add_paragraph('近期未检测到典型K线形态。')

        if summary:
            doc.add_paragraph('')
            doc.add_paragraph(summary[:250])
        if key_obs:
            doc.add_paragraph('')
            doc.add_paragraph(key_obs[:200])
    except Exception as e:
        doc.add_paragraph('K线形态分析暂不可用')
    # ═══ 四、量化评分 ═══
    add_colored_heading('四、量化评分（综合打分）', level=1)
    doc.add_paragraph('我们用7个因子给股票打分（满分100），综合考虑了涨跌幅、技术指标、基本面、成交量、风险等。')

    add_kpi_table(["评分维度", "得分", "满分", "说明"],
                  [[f"📈 动量", f"{qs_fs('momentum', quant):.0f}", "100", "近期涨幅大小。分太高=涨太多了(追高风险)，分太低=跌太多了(可能反弹)"],
                   [f"🔧 技术", f"{qs_fs('technical', quant):.0f}", "100", "MACD+RSI+KDJ等技术指标的综合判断"],
                   [f"💰 基本面", f"{qs_fs('fundamental', quant):.0f}", "100", "ROE/毛利率/营收增长等财务指标"],
                   [f"📊 量能", f"{qs_fs('volume', quant):.0f}", "100", "成交量和价格配合度。放量上涨好，缩量下跌差"],
                   [f"⚠️ 风险", f"{qs_fs('risk', quant):.0f}", "100", "波动率和回撤。分越低=风险越大"],
                   [f"🎯 综合", f"{qs:.0f}", "100", f"评级: {quant.get('rating', '')}。≥80强烈买入，≥60可以买，<40建议观望"]])

    doc.add_paragraph()
    doc.add_paragraph(f'短线评分: {st_sc:.0f} 分（{st.get("评级", "") if isinstance(st, dict) else ""}）  '
                      f'组合信号: {combo_sig}(强度{combo_str})  多周期共振: {mr.get("状态", "")}({mr_str})')
    doc.add_paragraph('👉 信号强度≥+4时买入信号强烈，0左右时方向不明，≤-4时建议卖出。')

    # ═══ 五、基本面 ═══
    add_colored_heading('五、基本面（公司本身怎么样）', level=1)
    doc.add_paragraph('基本面看的是公司本身的盈利能力和财务健康状况。')

    add_kpi_table(["指标", "数值", "说明"],
                  [[f"ROE(净资产收益率)", f"{roe:.2f}%", "巴菲特最看重的指标。>15%优秀，>10%良好，<5%偏弱。ROE高的公司更能赚钱"],
                   [f"PE(市盈率)", f"{pe}", "股价÷每股收益。PE越低越便宜。成长股PE较高，价值股PE较低"],
                   [f"PB(市净率)", f"{pb}", "股价÷每股净资产。PB<1可能被低估，银行股通常PB较低"],
                   [f"每股收益", f"{eps}", "每持有一股能分到多少利润。数值越高越好"],
                   [f"主力资金(5日)", f"{total_flow:+.2f}亿", "大资金最近5天是买入还是卖出。正数=大资金在买，负数=在卖"]])


    # Manipulator intention
    doc.add_page_break()
    doc.add_heading('庄家意图分析', level=1)
    doc.add_paragraph('识别庄家四阶段：建仓 → 洗盘 → 拉升 → 出货', style='Intense Quote')
    try:
        mi = analyze_manipulator_intention(kline, {'price': price, 'atr': atr})
        phase = mi.get('phase', '?')
        conf = mi.get('phase_confidence', 0)
        doc.add_paragraph('当前阶段：' + phase + '（置信度' + str(conf) + '%)')
        for s in mi.get('signals', []):
            doc.add_paragraph('  * ' + s)
        if mi.get('volume_analysis'):
            doc.add_paragraph('成交量分析：' + mi['volume_analysis'])
        if mi.get('assessment'):
            doc.add_paragraph('综合评估：' + mi['assessment'])
        if mi.get('risk_note'):
            doc.add_paragraph('风险提示：' + mi['risk_note'])
    except Exception as e:
        doc.add_paragraph('分析暂不可用：' + str(e)[:60])
    # ═══ 六、AI 预测 ═══
    add_colored_heading('六、AI 预测（机器学习判断）', level=1)
    doc.add_paragraph('AI用三个模型（XGBoost/RandomForest/LightGBM）分析历史数据，预测股票短期涨跌。')
    doc.add_paragraph(f'🤖 AI预测方向: {ml_dir}  置信度: {ml_conf:.0f}%  模型一致性: {ml_agree}')
    if ml_agree == "高":
        doc.add_paragraph('👉 三个AI模型意见一致，预测可信度较高。')
    else:
        doc.add_paragraph('👉 AI模型之间存在分歧，预测仅供参考，不能作为唯一依据。')

    # ═══ 七、多空辩论 ═══
    add_colored_heading('七、多空辩论（看涨 vs 看跌）', level=1)
    doc.add_paragraph('模拟多头（看涨方）和空头（看跌方）的辩论，帮你看清两边的理由。')

    p = doc.add_paragraph()
    p.add_run(f'🟢 看涨理由（{debate["bull"]["score"]}分）:').bold = True
    for pt in debate["bull"]["points"]:
        doc.add_paragraph(f'    ✅ {pt}')

    p = doc.add_paragraph()
    p.add_run(f'🔴 看跌理由（{debate["bear"]["score"]}分）:').bold = True
    for pt in debate["bear"]["points"]:
        doc.add_paragraph(f'    ❌ {pt}')

    if debate["bull"]["score"] > debate["bear"]["score"]:
        doc.add_paragraph(f'👉 结论: 看涨理由更充分 — {debate.get("action", "")}')
    elif debate["bull"]["score"] < debate["bear"]["score"]:
        doc.add_paragraph(f'👉 结论: 看跌理由更充分 — {debate.get("action", "")}')
    else:
        doc.add_paragraph(f'👉 结论: 双方势均力敌 — 建议观望等待方向明确')

    # ═══ 八、策略回测 ═══
    if bt:
        add_colored_heading('八、策略回测（历史表现如何）', level=1)
        doc.add_paragraph('用4种常见交易策略模拟这只股票的历史表现，找到最适合它的操作方式。')
        rows = []
        best_name = ""
        for s, r in bt.items():
            m = r['metrics']
            rows.append([r['name'], f"{m['总收益率%']:.0f}%", f"{m['夏普比率']:.2f}", f"-{m['最大回撤%']:.0f}%"])
            if m['夏普比率'] == max(x[1]['metrics']['夏普比率'] for x in bt.items()):
                best_name = r['name']
        add_kpi_table(["策略名称", "总收益率", "夏普比率(越高越好)", "最大回撤(越低越好)"], rows)
        doc.add_paragraph(f'👉 最优策略是「{best_name}」，历史上用这个策略操作这只股票效果最好。')

    # ═══ 九、操作建议 ═══
    add_colored_heading('九、操作建议（我该怎么做）', level=1)
    doc.add_paragraph('综合以上所有分析，给出具体的买入/卖出建议。')

    action_color = "22c55e" if buy_signals >= 4 else ("f59e0b" if buy_signals >= 2 else "ef4444")
    action_text = "可以买入" if buy_signals >= 4 else ("轻仓试探" if buy_signals >= 2 else "建议观望或减仓")

    p = doc.add_paragraph()
    pr = p.add_run(f'📌 {action_text}'); pr.font.size = Pt(16); pr.bold = True
    pr.font.color.rgb = RGBColor(int(action_color[0:2], 16), int(action_color[2:4], 16), int(action_color[4:6], 16))

    add_kpi_table(["操作", "价格", "说明"],
                  [[f"🟢 买入区间", f"{buy_low} - {buy_high} 元", f"在这个价格范围内买入比较安全。当前价{price}元{'在区间内，可以考虑' if buy_low <= price <= buy_high else '，等回调到区间再买'}"],
                   [f"🔴 止损价", f"{stop.get('止损参考价', price * 0.93):.2f} 元", f"如果跌破这个价格，说明判断错了，必须卖出止损。亏小钱比亏大钱好"],
                   [f"🎯 止盈价", f"{stop.get('止盈参考价', price * 1.07):.2f} 元", f"涨到这个价格附近可以考虑卖出获利。不要贪心，有赚就好"],
                   [f"⏱️ 建议持有", "3-5 个交易日", "短线操作不要太久，快进快出"]])

    doc.add_paragraph()
    doc.add_paragraph(f'💡 小贴士：'
                      f'止损价是最后的底线，到了必须走；止盈价是参考目标，到了可以考虑卖。'
                      f'如果买入后3天还没涨，也应该考虑退出。')

    # ═══ 十、宏观环境 ═══
    if macro and 'error' not in macro:
        add_colored_heading('十、宏观环境', level=1)
        ind = macro.get('数据', {})
        doc.add_paragraph(f'PMI: {ind.get("制造业PMI", "?")}（>50经济扩张，<50经济收缩）  '
                          f'M2: {ind.get("M2同比%", "?")}%（货币供应量增速，越高钱越多）')
        for sig in macro.get('信号', []):
            doc.add_paragraph(f'  • {sig}')
        doc.add_paragraph(f'整体判断: {macro.get("整体", "")}')

    # ═══ 免责 ═══
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    p = doc.add_paragraph()
    p.add_run('⚠️ 免责声明：').bold = True
    p.add_run('本报告由AI自动生成，所有分析仅供参考学习，不构成任何投资建议。'
              '股市有风险，投资需谨慎。请根据自身情况独立做出投资决策。')
    doc.add_paragraph(f'数据来源: 新浪财经(实时行情) | Baostock(行业分类) | akshare(基本面) | 东方财富(板块资金)')
    doc.add_paragraph(f'数据时间: K线来源 {str(kline.iloc[-1][chr(26085)+chr(26399)])[:10]}  |  实时行情 {datetime.now().strftime("%H:%M")} (非交易日为上一交易日)')

    path = f'reports/{name}_深度分析_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    os.makedirs('reports', exist_ok=True)
    doc.save(path)
    print(f'报告已生成: {path}')
    return path


def qs_fs(key, quant):
    fs = quant.get('factor_scores', {})
    v = fs.get(key, {})
    return float(v.get('score', 0)) if isinstance(v, dict) else 0


if __name__ == '__main__':
    code = sys.argv[1] if len(sys.argv) > 1 else '600066'
    generate(code)
