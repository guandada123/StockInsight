"""DOCX 报告生成模块"""

import os
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .config import REPORT_DIR


def _set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _make_header_row(table, headers, color="4472C4"):
    """设置表头行"""
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, color)


def _add_table_row(table, values, bold=False):
    """添加数据行"""
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(v))
        run.font.size = Pt(9)
        if bold:
            run.bold = True


def _create_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _make_header_row(table, headers)
    for row_vals in rows:
        _add_table_row(table, row_vals)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


def add_title(doc, text, level=0):
    """添加标题"""
    if level == 0:
        p = doc.add_heading(text, level=1)
    else:
        p = doc.add_heading(text, level=level + 1)
    return p


def add_para(doc, text, bold=False, size=10, color=None, align=None):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_image(doc, path, width=14):
    """添加图片，存在时"""
    if path and os.path.exists(path):
        doc.add_picture(path, width=Cm(width))
        return True
    return False


def generate_report(sectors_df, picks_df, stock_details, output_name=None):
    """生成完整DOCX报告"""
    os.makedirs(REPORT_DIR, exist_ok=True)
    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")
    output_name = output_name or f"股票分析报告_{date_str}.docx"
    output_path = os.path.join(REPORT_DIR, output_name)

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ══════════ 封面 ══════════
    for _ in range(6):
        doc.add_paragraph()

    add_title(doc, "股票市场综合分析报告", level=0)
    add_para(doc, f"报告日期：{date_str}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc, f"生成时间：{now.strftime('%Y-%m-%d %H:%M')}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER
    )
    add_para(
        doc,
        "数据来源：东方财富 / 同花顺",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x88, 0x88, 0x88),
    )

    doc.add_page_break()

    # ══════════ 目录页 ══════════
    add_title(doc, "目录", level=0)
    toc_items = [
        "一、市场板块排行榜",
        "二、TOP 3 板块深度分析",
        "    2.1 板块概况",
        "    2.2 推荐个股技术分析",
        "    2.3 资金流向分析",
        "    2.4 基本面评分",
        "三、推荐股票汇总",
        "四、免责声明",
    ]
    for item in toc_items:
        add_para(doc, item, size=12)

    doc.add_page_break()

    # ══════════ 一、板块排行榜 ══════════
    add_title(doc, "一、市场板块排行榜", level=0)
    add_para(doc, f"截至 {now.strftime('%Y-%m-%d %H:%M')}，全市场行业板块排名如下：", size=10)

    if not sectors_df.empty:
        top15 = sectors_df.head(15)
        headers = ["排名", "板块名称", "涨跌幅", "上涨/总数", "成交额(亿)", "评分"]
        rows = []
        for i, (_, r) in enumerate(top15.iterrows(), 1):
            total = int(r["上涨家数"] + r["下跌家数"])
            vol = r["总成交额"] / 1e8 if r["总成交额"] else 0
            rows.append(
                [
                    str(i),
                    r["板块名称"],
                    f"{r['涨跌幅']:+.2f}%",
                    f"{int(r['上涨家数'])}/{total}",
                    f"{vol:.1f}",
                    f"{r['评分']:.2f}",
                ]
            )
        _create_table(doc, headers, rows)

    doc.add_page_break()

    # ══════════ 二、TOP 3 板块深度分析 ══════════
    add_title(doc, "二、TOP 3 板块深度分析", level=0)

    top3 = sectors_df.head(3) if not sectors_df.empty else pd.DataFrame()
    for rank, (_, sec_row) in enumerate(top3.iterrows(), 1):
        sec_name = sec_row["板块名称"]
        total = int(sec_row["上涨家数"] + sec_row["下跌家数"])
        vol = sec_row["总成交额"] / 1e8

        add_title(doc, f"2.{rank} {sec_name}", level=1)

        # 2.x.1 板块概况
        add_title(doc, "板块概况", level=2)
        _create_table(
            doc,
            ["指标", "数值"],
            [
                ["涨跌幅", f"{sec_row['涨跌幅']:+.2f}%"],
                ["上涨/下跌", f"{int(sec_row['上涨家数'])} / {int(sec_row['下跌家数'])}"],
                ["成交额", f"{vol:.1f}亿"],
                ["评分", f"{sec_row['评分']:.2f}"],
            ],
        )

        # 该板块的推荐个股
        sector_picks = (
            picks_df[picks_df["板块"] == sec_name] if not picks_df.empty else pd.DataFrame()
        )
        for _, stk in sector_picks.iterrows():
            code = stk["代码"]
            name = stk["名称"]
            add_title(doc, f"推荐个股：{code} {name}", level=2)

            detail = stock_details.get(code, {})

            # 技术图表
            kline_path = detail.get("kline_chart", "")
            ind_path = detail.get("indicator_chart", "")
            flow_path = detail.get("fundflow_chart", "")

            tech = detail.get("technical", {})

            if kline_path and os.path.exists(kline_path):
                success = add_image(doc, kline_path, width=15)
                if success:
                    add_para(
                        doc,
                        f"图：{code} {name} K线及MACD图",
                        size=9,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        color=RGBColor(0x88, 0x88, 0x88),
                    )

            if ind_path and os.path.exists(ind_path):
                success = add_image(doc, ind_path, width=15)
                if success:
                    add_para(
                        doc,
                        f"图：{code} {name} RSI及KDJ指标",
                        size=9,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        color=RGBColor(0x88, 0x88, 0x88),
                    )

            # 技术指标汇总
            if tech:
                add_para(doc, "技术指标分析：", bold=True, size=11)
                add_para(
                    doc,
                    f"最新收盘价：{tech.get('最新收盘', '-')}元  |  "
                    f"当日涨跌幅：{tech.get('涨跌幅', 0):+.2f}%  |  "
                    f"近5日涨幅：{tech.get('近5日涨跌幅', '-')}%  |  "
                    f"近20日涨幅：{tech.get('近20日涨跌幅', '-')}%",
                    size=10,
                )

                macd = tech.get("MACD", {})
                rsi = tech.get("RSI", {})
                kdj = tech.get("KDJ", {})

                _create_table(
                    doc,
                    ["指标", "数值", "信号"],
                    [
                        [
                            "MACD",
                            f"DIF={macd.get('DIF', '-')} DEA={macd.get('DEA', '-')}",
                            macd.get("信号", "-"),
                        ],
                        ["RSI", str(rsi.get("值", "-")), rsi.get("信号", "-")],
                        [
                            "KDJ",
                            f"K={kdj.get('K', '-')} D={kdj.get('D', '-')} J={kdj.get('J', '-')}",
                            kdj.get("信号", "-"),
                        ],
                    ],
                )

                # 均线
                ma_info = tech.get("均线", {})
                if ma_info:
                    add_para(doc, "均线位置：", bold=True, size=10)
                    ma_lines = []
                    for w in [5, 10, 20, 60]:
                        m = ma_info.get(f"MA{w}", {})
                        if m:
                            ma_lines.append(f"MA{w}={m['值']}（股价在{m['股价位置']}）")
                    add_para(doc, "  ".join(ma_lines), size=9)

            # 资金流向图
            if flow_path and os.path.exists(flow_path):
                success = add_image(doc, flow_path, width=14)
                if success:
                    add_para(
                        doc,
                        f"图：{code} {name} 近20日资金流向",
                        size=9,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        color=RGBColor(0x88, 0x88, 0x88),
                    )
            else:
                add_para(doc, "资金流向数据暂未获取到", size=9, color=RGBColor(0x88, 0x88, 0x88))

            # 基本面评分
            funda = detail.get("fundamentals", {})
            funda_score = detail.get("fundamental_score")
            if funda or funda_score is not None:
                add_para(doc, "基本面评分：", bold=True, size=11)
                _create_table(
                    doc,
                    ["指标", "数值"],
                    [
                        [k, f"{v:.2f}" if isinstance(v, float) else str(v)]
                        for k, v in funda.items()
                        if v is not None
                    ],
                )
                add_para(
                    doc,
                    f"综合基本面评分：{funda_score}/100",
                    bold=True,
                    size=11,
                    color=RGBColor(0x00, 0x80, 0x00)
                    if funda_score and funda_score >= 60
                    else RGBColor(0xCC, 0x00, 0x00),
                )

            doc.add_page_break()

    # ══════════ 三、推荐汇总 ══════════
    add_title(doc, "三、推荐股票汇总", level=0)

    if not picks_df.empty:
        headers = ["代码", "名称", "板块", "现价", "涨幅", "量比", "基本评分", "技术评分"]
        rows = []
        for _, stk in picks_df.iterrows():
            code = stk["代码"]
            detail = stock_details.get(code, {})
            fs = detail.get("fundamental_score", "-")
            ts = detail.get("tech_score", "-")
            rows.append(
                [
                    code,
                    stk["名称"],
                    stk["板块"],
                    f"{stk['最新价']:.2f}" if stk["最新价"] else "-",
                    f"{stk['涨跌幅']:+.2f}%",
                    f"{stk['量比']:.2f}",
                    f"{fs}/100" if fs != "-" else "-",
                    f"{ts}/100" if ts != "-" else "-",
                ]
            )
        _create_table(doc, headers, rows)

        add_para(doc, "", size=6)
        add_para(
            doc,
            "说明：基本评分基于ROE/营收增长/毛利率等财务指标，技术评分基于均线/MACD/RSI/KDJ等技术指标。"
            "分数越高代表该维度表现越好。",
            size=9,
            color=RGBColor(0x88, 0x88, 0x88),
        )

    doc.add_page_break()

    # ══════════ 四、免责声明 ══════════
    add_title(doc, "四、免责声明", level=0)
    disclaimers = [
        "本报告由AI自动生成，数据来源于东方财富、同花顺等公开市场数据。",
        "本报告仅供参考，不构成任何投资建议或投资承诺。",
        "投资有风险，入市需谨慎。过往表现不代表未来收益。",
        "报告生成时间：" + now.strftime("%Y-%m-%d %H:%M"),
        "AI工具：Claude Code",
    ]
    for d in disclaimers:
        add_para(doc, d, size=10, color=RGBColor(0x66, 0x66, 0x66))

    doc.save(output_path)
    return output_path
